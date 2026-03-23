# -*- coding: utf-8 -*-
from PyQt5 import QtCore, QtGui, QtWidgets
from main import generate_dual_graph
import os
from main import place_names
from docx import Document
from docx.shared import Inches
from worker import Worker
from docx.shared import Inches
from trend_analyzer import generate_trend_graphs
def replace_text_placeholder(doc, placeholder, value):

    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, str(value))

def replace_placeholder_with_image(doc, placeholder, image_stream, width=5):

    for paragraph in doc.paragraphs:

        if placeholder in paragraph.text:

            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, "")
                    run.add_picture(image_stream, width=Inches(width))

def replace_placeholder_with_image_grid(doc, placeholder, images, width=3):

    for paragraph in doc.paragraphs:

        if placeholder in paragraph.text:

            paragraph.text = paragraph.text.replace(placeholder, "")

            table = paragraph._element.addnext(
                doc.add_table(rows=2, cols=2)._element
            )

            table = paragraph._element.getnext()

            idx = 0
            table_obj = doc.tables[-1]

            for r in range(2):
                for c in range(2):

                    if idx < len(images):

                        cell = table_obj.rows[r].cells[c]
                        cell.paragraphs[0].add_run().add_picture(
                            images[idx],
                            width=Inches(width)
                        )

                    idx += 1

            break
STATIC_SCORES = {
    "Monasterio de Nuestra Senhora del Risco": 54,
    "Santuario de Nuestra Senhora de Sonsoles": 76,
    "Parque de El Soto Avila": 78,
    "Dolmen del Prado de Las Cruces": 67,
    "Ermita de San Segundo Avila": 78,
    "Iglesia de San Pedro Apostol Bernuy": 63,
    "Iglesia de Nuestra Sra del Rosario Vizolozano": 63,
    "Castro de las Cogotas": 80,
    "Arco de conejeros": 61,
    "Las Tres Cruces Cardenhosa": 65,
    "Iglesia de la Invencion de la Santa Cruz Cardenhosa": 71,
    "Castro de la Mesa de Miranda": 71,
    "Los Henrenes Cillan": 68,
    "Ermita San Miguel La Hija de Dios": 67,
    "Necropolis de Oco": 63,
    "Castillo de Manqueospese": 58,
    "Las Piedras de Garoza": 64,
    "Ermita Nuestra Senhora de las Fuentes": 78,
    "Necropolis de la Coba": 64,
    "Castro-Ulaca": 85,
    "Castillo de Villaviciosa Solosancho": 73,
    "Iglesia de la Transfiguracion del Senhor Vadillo": 74,
    "Verraco Campillo": 73,
    "Iglesia de La Natividad De Nuestra Senhora Campillo": 70
}


class Ui_Dialog(object):

    def setupUi(self, Dialog):
        Dialog.setObjectName("Dialog")
        Dialog.resize(400, 310)

        self.CalculateBotton = QtWidgets.QPushButton(Dialog)
        self.CalculateBotton.setGeometry(QtCore.QRect(150, 265, 100, 30))

        self.progressBar = QtWidgets.QProgressBar(Dialog)
        self.progressBar.setGeometry(QtCore.QRect(20, 230, 371, 23))

        self.SelectPlace = QtWidgets.QComboBox(Dialog)
        self.SelectPlace.setGeometry(QtCore.QRect(20, 40, 351, 22))

        self.place_label = QtWidgets.QLabel(Dialog)
        self.place_label.setGeometry(QtCore.QRect(20, 20, 200, 13))

        self.BrowseImportPath = QtWidgets.QLineEdit(Dialog)
        self.BrowseImportPath.setGeometry(QtCore.QRect(20, 100, 351, 27))

        self.BrowseOutputPath = QtWidgets.QLineEdit(Dialog)
        self.BrowseOutputPath.setGeometry(QtCore.QRect(20, 160, 351, 27))

        self.ImportPath = QtWidgets.QLabel(Dialog)
        self.ImportPath.setGeometry(QtCore.QRect(20, 80, 200, 16))

        self.OutputPath = QtWidgets.QLabel(Dialog)
        self.OutputPath.setGeometry(QtCore.QRect(20, 140, 200, 16))

        self.FinalScore = QtWidgets.QLabel(Dialog)
        self.FinalScore.setGeometry(QtCore.QRect(20, 200, 200, 20))

        font = QtGui.QFont()
        font.setBold(True)
        font.setPointSize(10)
        self.FinalScore.setFont(font)

        self.retranslateUi(Dialog)
        QtCore.QMetaObject.connectSlotsByName(Dialog)

    def retranslateUi(self, Dialog):
        _translate = QtCore.QCoreApplication.translate
        Dialog.setWindowTitle(_translate("Dialog", "Analizador de Reseñas"))
        self.CalculateBotton.setText(_translate("Dialog", "Calculate"))
        self.place_label.setText(_translate("Dialog", "Place"))
        self.ImportPath.setText(_translate("Dialog", "Experts result path"))
        self.OutputPath.setText(_translate("Dialog", "Output path"))
        self.FinalScore.setText(_translate("Dialog", "Final score: --"))


class MainWindow(QtWidgets.QDialog):

    def __init__(self):
        super().__init__()

        self.ui = Ui_Dialog()
        self.ui.setupUi(self)

        for place_id, place_name in place_names.items():
            display_name = place_name.replace("_", " ")
            self.ui.SelectPlace.addItem(display_name, place_id)

        self.ui.CalculateBotton.clicked.connect(self.calculate)

        self.ui.ExpertButton = QtWidgets.QPushButton("...", self)
        self.ui.ExpertButton.setGeometry(QtCore.QRect(375, 100, 25, 27))
        self.ui.ExpertButton.clicked.connect(self.select_expert_file)

        self.ui.OutputButton = QtWidgets.QPushButton("...", self)
        self.ui.OutputButton.setGeometry(QtCore.QRect(375, 160, 25, 27))
        self.ui.OutputButton.clicked.connect(self.select_output_folder)

    def calculate(self):

        place_id = self.ui.SelectPlace.currentData()
        output_folder = self.ui.BrowseOutputPath.text().strip()
        if not output_folder:
                    QtWidgets.QMessageBox.warning(
                        self,
                        "Ruta incompleta",
                        "Por favor, selecciona una carpeta de salida antes de calcular."
                    )
                    return
        self.ui.progressBar.setValue(0)
        self.ui.FinalScore.setText("Final score: Calculating...")

        self.worker = Worker(place_id, output_folder)
        self.worker.progress.connect(self.ui.progressBar.setValue)
        self.worker.finished.connect(self.on_finished)
        self.worker.start()

    def select_expert_file(self):

        file_path, _ = QtWidgets.QFileDialog.getOpenFileName(
            self,
            "Select Word template",
            "",
            "Word Files (*.docx);;All Files (*)"
        )

        if file_path:
            self.ui.BrowseImportPath.setText(file_path)

    def select_output_folder(self):

        folder_path = QtWidgets.QFileDialog.getExistingDirectory(
            self,
            "Select output folder"
        )

        if folder_path:
            self.ui.BrowseOutputPath.setText(folder_path)

    def on_finished(self, result):

        final_score_user = result["final_score"]

        place_name = self.ui.SelectPlace.currentText().strip().lower()

        static_score = None
        for key, value in STATIC_SCORES.items():
            if key.lower() == place_name:
                static_score = value
                break
        
       

        if static_score:
            static_score = static_score / 20  
            final_score = (final_score_user + static_score) / 2
          
        else:
            final_score = final_score_user
        self.ui.FinalScore.setText(f"Final score: {round(final_score,2)}")
        self.ui.progressBar.setValue(100)

        output_folder = self.worker.output_path
        template_path = self.ui.BrowseImportPath.text().strip()

        image_path = os.path.join(output_folder, "double_graph.png")
        total_reviews = result["total_reviews"]
        rating_score = round(result["final_score"], 1)
        generate_dual_graph(
            x_data=result["x_data"],
            y_rating=result["y_rating"],
            y_votes=result["y_votes"],
            output_path=image_path
        )

        try:

            doc = Document(template_path)

            double_graph = os.path.join(output_folder, "double_graph.png")

            rating_images, mean_images, ratings_time_images = generate_trend_graphs(
                self.worker.place_id,
                output_folder
            )
            
            replace_text_placeholder(doc, "{{REVIEWS}}", total_reviews)
            replace_text_placeholder(doc, "{{NOTE}}", rating_score)
            replace_placeholder_with_image(doc, "{{DOUBLE_GRAPH}}", double_graph, width=6)

            replace_placeholder_with_image_grid(
                doc,
                "{{RATING_EVOLUTION_GRID}}",
                rating_images
            )

            replace_placeholder_with_image_grid(
                doc,
                "{{MEAN_EVOLUTION_GRID}}",
                mean_images
            )

            replace_placeholder_with_image_grid(
                doc,
                "{{RATINGS_OVER_TIME_GRID}}",
                ratings_time_images
            )

            output_file = os.path.join(output_folder, "resultado_final.docx")
            doc.save(output_file)
        except Exception as e:

            QtWidgets.QMessageBox.critical(
                self,
                "Error",
                f"Ocurrió un error:\n{str(e)}"
            )

if __name__ == "__main__":

    import sys

    app = QtWidgets.QApplication(sys.argv)
    window = MainWindow()
    window.show()

    sys.exit(app.exec_())